from flask import Flask, request, jsonify
from flask_cors import CORS
import pdfplumber
from docx import Document
import spacy
from openai import OpenAI
from dotenv import load_dotenv
import os
from collections import Counter
import re
import io
import logging

# Logging einrichten
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

# Lade Umgebungsvariablen
load_dotenv()

# NLP-Modell laden
try:
    nlp = spacy.load("de_core_news_md")
except Exception as e:
    logger.error(f"Fehler beim Laden des spaCy-Modells: {e}")
    exit(1)

# OpenAI API initialisieren
try:
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    if not os.getenv("OPENAI_API_KEY"):
        logger.error("OpenAI API-Schlüssel nicht gefunden.")
        exit(1)
except Exception as e:
    logger.error(f"Fehler beim Initialisieren der OpenAI API: {e}")
    exit(1)

# Funktionen für Textextraktion
def extract_text_from_pdf(file):
    try:
        logger.debug(f"Verarbeite PDF-Datei: {file.filename}")
        with pdfplumber.open(file) as pdf:
            text = "".join(page.extract_text() or "" for page in pdf.pages)
        if not text.strip():
            logger.warning("Kein Text in der PDF-Datei gefunden.")
            return "Fehler: Kein extrahierbarer Text in der PDF-Datei."
        logger.debug(f"Extrahierter Text (ersten 100 Zeichen): {text[:100]}")
        return text
    except Exception as e:
        logger.error(f"Fehler beim Lesen der PDF: {e}")
        return f"Fehler beim Lesen der PDF: {str(e)}"

def extract_text_from_docx(file):
    try:
        logger.debug(f"Verarbeite DOCX-Datei: {file.filename}")
        doc = Document(file)
        text = "".join(paragraph.text for paragraph in doc.paragraphs)
        if not text.strip():
            logger.warning("Kein Text in der DOCX-Datei gefunden.")
            return "Fehler: Kein extrahierbarer Text in der DOCX-Datei."
        logger.debug(f"Extrahierter Text (ersten 100 Zeichen): {text[:100]}")
        return text
    except Exception as e:
        logger.error(f"Fehler beim Lesen der DOCX: {e}")
        return f"Fehler beim Lesen der DOCX: {str(e)}"

# Verbesserte Schlüsselwort-Extraktion
blacklisted_terms = [
    "daten", "persönliche", "nachname", "geburtsort", "geburtstag",
    "adresse", "telefon", "email", "e-mail", "straße", "postleitzahl"
]
priority_terms = [
    "vermögensberatung", "kundenberatung", "finanzwesen", "bankkauffrau",
    "bankkaufmann", "kundenmanagement", "finanzprodukte"
]

def extract_keywords(text):
    doc = nlp(text)
    keywords = [
        token.text.lower() for token in doc 
        if token.pos_ in ["NOUN", "ADJ", "VERB"] 
        and len(token.text) > 3 
        and not token.is_stop
        and token.text.lower() not in blacklisted_terms
    ]
    keyword_counts = Counter(keywords)
    prioritized = [word for word in priority_terms if word in keywords]
    other_keywords = [word for word, count in keyword_counts.most_common(5 - len(prioritized))]
    return prioritized + other_keywords[:5 - len(prioritized)]

# Eingabevalidierung
def is_valid_input(text):
    return len(text) >= 10 and bool(re.search(r'[a-zA-ZäöüÄÖÜß]{3,}', text))

# Funktion für Motivationsschreiben
def generate_cover_letter(cv_text, job_text, strengths, weaknesses, name, company, tone="Professionell"):
    if not all([cv_text, job_text, strengths, weaknesses, name, company]):
        return "Fehler: Alle Felder müssen ausgefüllt sein."
    if not is_valid_input(strengths) or not is_valid_input(weaknesses):
        return "Fehler: Stärken und Schwächen müssen sinnvolle Eingaben sein."

    cv_keywords = extract_keywords(cv_text)
    job_keywords = extract_keywords(job_text)
    
    prompt = f"""
    Erstelle ein {tone.lower()}es Motivationsschreiben auf Deutsch für {name}, der/die sich bei {company} auf eine Position als Bankkauffrau/-mann bewirbt. Verwende die folgenden Informationen:

    **Lebenslauf (Auszug)**: {cv_text[:1000]}
    **Wichtige Schlüsselwörter aus dem Lebenslauf**: {', '.join(cv_keywords)}
    **Stellenprofil (Auszug)**: {job_text[:1000]}
    **Wichtige Schlüsselwörter aus dem Stellenprofil**: {', '.join(job_keywords)}
    **Stärken**: {strengths}
    **Schwächen**: {weaknesses}

    Das Schreiben soll:
    - Höflich und professionell sein.
    - Die Qualifikationen des Bewerbers mit den Anforderungen der Stelle verknüpfen.
    - Stärken hervorheben und Schwächen positiv darstellen.
    - Maximal 400 Wörter lang sein.
    - Mit einer höflichen Schlussformel enden.
    """
    
    try:
        logger.debug("Sende Anfrage an OpenAI API...")
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Du bist ein professioneller Bewerbungsassistent."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=600,
            temperature=0.7
        )
        logger.debug("OpenAI API-Antwort erhalten.")
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"Fehler bei der OpenAI API-Anfrage: {e}")
        return f"Fehler bei der API-Anfrage: {str(e)}"

# API-Endpunkt
@app.route('/generate_cover_letter', methods=['POST'])
def api_generate_cover_letter():
    logger.debug("Erhalte POST-Anfrage für /generate_cover_letter")
    logger.debug(f"Formulardaten: {request.form.to_dict()}")
    logger.debug(f"Dateien: {list(request.files.keys())}")
    logger.debug(f"Content-Type: {request.content_type}")
    try:
        # Dateien und Formulardaten verarbeiten
        cv_file = request.files.get('cv_file')
        job_file = request.files.get('job_file')
        strengths = request.form.get('strengths')
        weaknesses = request.form.get('weaknesses')
        name = request.form.get('name')
        company = request.form.get('company')
        tone = request.form.get('tone', 'Professionell')

        # Validierung
        if not cv_file:
            logger.error("Kein Lebenslauf hochgeladen.")
            return jsonify({"error": "Bitte laden Sie einen Lebenslauf hoch."}), 400
        if not job_file:
            logger.error("Kein Stellenprofil hochgeladen.")
            return jsonify({"error": "Bitte laden Sie ein Stellenprofil hoch."}), 400

        # Überprüfe MIME-Type und Dateigröße
        logger.debug(f"Lebenslauf-Datei: {cv_file.filename}, MIME-Type: {cv_file.mimetype}, Größe: {len(cv_file.read())} Bytes")
        cv_file.seek(0)  # Zurück zum Anfang der Datei
        logger.debug(f"Stellenprofil-Datei: {job_file.filename}, MIME-Type: {job_file.mimetype}, Größe: {len(job_file.read())} Bytes")
        job_file.seek(0)

        # Textextraktion
        cv_text = ""
        if cv_file.mimetype == 'application/pdf':
            cv_text = extract_text_from_pdf(cv_file)
        elif cv_file.mimetype == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            cv_text = extract_text_from_docx(cv_file)
        else:
            logger.error(f"Ungültiges Lebenslauf-Format: {cv_file.mimetype}")
            return jsonify({"error": f"Ungültiges Lebenslauf-Format: {cv_file.mimetype}. Nur PDF oder DOCX erlaubt."}), 400

        job_text = ""
        if job_file.mimetype == 'application/pdf':
            job_text = extract_text_from_pdf(job_file)
        else:
            logger.error(f"Ungültiges Stellenprofil-Format: {job_file.mimetype}")
            return jsonify({"error": f"Ungültiges Stellenprofil-Format: {job_file.mimetype}. Nur PDF erlaubt."}), 400

        # Validierung
        if "Fehler" in cv_text:
            logger.error(f"Fehler beim Verarbeiten des Lebenslaufs: {cv_text}")
            return jsonify({"error": f"Fehler beim Verarbeiten des Lebenslaufs: {cv_text}"}), 400
        if "Fehler" in job_text:
            logger.error(f"Fehler beim Verarbeiten des Stellenprofils: {job_text}")
            return jsonify({"error": f"Fehler beim Verarbeiten des Stellenprofils: {job_text}"}), 400

        if not all([name, company, strengths, weaknesses]):
            logger.error("Fehlende Textfelder in der Anfrage.")
            return jsonify({"error": "Alle Textfelder müssen ausgefüllt sein."}), 400
        if not is_valid_input(strengths) or not is_valid_input(weaknesses):
            logger.error("Ungültige Stärken oder Schwächen.")
            return jsonify({"error": "Stärken und Schwächen müssen sinnvolle Eingaben sein."}), 400

        # Motivationsschreiben generieren
        logger.debug("Generiere Motivationsschreiben...")
        cover_letter = generate_cover_letter(cv_text, job_text, strengths, weaknesses, name, company, tone)

        if "Fehler" in cover_letter:
            logger.error(f"Fehler beim Generieren des Motivationsschreibens: {cover_letter}")
            return jsonify({"error": cover_letter}), 500

        logger.debug("Motivationsschreiben erfolgreich generiert.")
        return jsonify({
            "cover_letter": cover_letter,
            "cv_keywords": extract_keywords(cv_text),
            "job_keywords": extract_keywords(job_text)
        })
    except Exception as e:
        logger.error(f"Serverfehler: {str(e)}")
        return jsonify({"error": str(e)}), 500

# API-Endpunkt für DOCX
from flask import send_file

@app.route('/generate_docx', methods=['POST'])
def generate_docx():
    try:
        logger.debug("Neue Anfrage an /generate_docx")
        data = request.get_json()
        cover_letter = data.get('cover_letter')
        if not cover_letter:
            logger.error("Kein Motivationsschreiben angegeben.")
            return jsonify({"error": "Kein Motivationsschreiben angegeben."}), 400
        
        doc = Document()
        doc.add_heading("Motivationsschreiben", level=1)
        for paragraph in cover_letter.split("\n\n"):
            doc.add_paragraph(paragraph.strip())
        
        filename = "Motivationsschreiben.docx"
        doc.save(filename)
        logger.debug(f"DOCX-Datei erstellt: {filename}")
        return send_file(filename, as_attachment=True)
    except Exception as e:
        logger.error(f"Fehler beim Generieren der DOCX: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0', port=5000)