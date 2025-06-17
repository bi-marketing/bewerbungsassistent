import streamlit as st
import pdfplumber
from docx import Document
import spacy
from openai import OpenAI
from dotenv import load_dotenv
import os
from collections import Counter
import re

# Lade Umgebungsvariablen aus .env
load_dotenv()

# NLP-Modell laden
try:
    nlp = spacy.load("de_core_news_md")
except Exception as e:
    st.error(f"Fehler beim Laden des spaCy-Modells: {e}")
    st.stop()

# OpenAI API initialisieren
try:
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    if not os.getenv("OPENAI_API_KEY"):
        st.error("OpenAI API-Schlüssel nicht gefunden. Bitte stelle sicher, dass er in der .env-Datei korrekt eingetragen ist.")
        st.stop()
except Exception as e:
    st.error(f"Fehler beim Initialisieren der OpenAI API: {e}")
    st.stop()

# Funktionen für Textextraktion
def extract_text_from_pdf(file):
    try:
        with pdfplumber.open(file) as pdf:
            text = "".join(page.extract_text() or "" for page in pdf.pages)
        return text
    except Exception as e:
        return f"Fehler beim Lesen der PDF: {e}"

def extract_text_from_docx(file):
    try:
        doc = Document(file)
        text = "".join(paragraph.text for paragraph in doc.paragraphs)
        return text
    except Exception as e:
        return f"Fehler beim Lesen der DOCX: {e}"

# Verbesserte Schlüsselwort-Extraktion
blacklisted_terms = [
    "daten", "persönliche", "nachname", "geburtsort", "geburtstag",
    "adresse", "telefon", "email", "e-mail", "straße", "postleitzahl"
]
priority_terms = [
    "vermögensberatung", "kundenberatung", "finanzwesen", "bankkauffrau",
    "bankkaufmann", "kundenmanagement", "finanzprodukte"
]

def extract_keywords(text):
    doc = nlp(text)
    keywords = [
        token.text.lower() for token in doc 
        if token.pos_ in ["NOUN", "ADJ", "VERB"] 
        and len(token.text) > 3 
        and not token.is_stop
        and token.text.lower() not in blacklisted_terms
    ]
    keyword_counts = Counter(keywords)
    prioritized = [word for word in priority_terms if word in keywords]
    other_keywords = [word for word, count in keyword_counts.most_common(5 - len(prioritized))]
    return prioritized + other_keywords[:5 - len(prioritized)]

# Eingabevalidierung
def is_valid_input(text):
    return len(text) >= 10 and bool(re.search(r'[a-zA-ZäöüÄÖÜß]{3,}', text))

# Funktion für Motivationsschreiben mit OpenAI API
def generate_cover_letter(cv_text, job_text, strengths, weaknesses, name, company):
    cv_keywords = extract_keywords(cv_text)
    job_keywords = extract_keywords(job_text)
    
    prompt = f"""
    Erstelle ein professionelles, flüssiges und individuelles Motivationsschreiben auf Deutsch für {name}, der/die sich bei {company} auf eine Position als Bankkauffrau/-mann bewirbt. Verwende die folgenden Informationen:

    **Lebenslauf (Auszug)**: {cv_text[:1000]}
    **Wichtige Schlüsselwörter aus dem Lebenslauf**: {', '.join(cv_keywords)}
    **Stellenprofil (Auszug)**: {job_text[:1000]}
    **Wichtige Schlüsselwörter aus dem Stellenprofil**: {', '.join(job_keywords)}
    **Stärken**: {strengths}
    **Schwächen**: {weaknesses}

    Das Schreiben soll:
    - Höflich und professionell sein.
    - Die Qualifikationen des Bewerbers mit den Anforderungen der Stelle verknüpfen.
    - Stärken hervorheben und Schwächen positiv darstellen.
    - Maximal 400 Wörter lang sein.
    - Mit einer höflichen Schlussformel enden.
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",  # Oder "gpt-4" wenn verfügbar
            messages=[
                {"role": "system", "content": "Du bist ein professioneller Bewerbungsassistent."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=600,
            temperature=0.7
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Fehler bei der API-Anfrage: {e}"

# Funktion zum Speichern als DOCX
def save_cover_letter(text, filename="Motivationsschreiben.docx"):
    doc = Document()
    doc.add_heading("Motivationsschreiben", level=1)
    for paragraph in text.split("\n\n"):
        doc.add_paragraph(paragraph.strip())
    doc.save(filename)
    return filename

# Streamlit-App
st.title("Bewerbungsassistent")
st.write("Bitte laden Sie Ihren Lebenslauf und das Stellenprofil hoch und geben Sie Ihre Stärken und Schwächen ein.")

# Eingabefelder
name = st.text_input("Ihr Name", value="Max Mustermann")
company = st.text_input("Name des Unternehmens", value="Beispiel GmbH")
cv_file = st.file_uploader("Lebenslauf hochladen (PDF oder DOCX)", type=["pdf", "docx"])
job_file = st.file_uploader("Stellenprofil hochladen (PDF)", type=["pdf"])
strengths = st.text_area("Ihre Stärken (z. B. Teamfähigkeit, Zuverlässigkeit)")
weaknesses = st.text_area("Ihre Schwächen (z. B. Perfektionismus)")
tone = st.selectbox("Ton des Schreibens", ["Professionell", "Kreativ", "Formell"])

# Button zum Generieren
if st.button("Motivationsschreiben generieren"):
    # Validierung
    if not is_valid_input(strengths) or not is_valid_input(weaknesses):
        st.error("Bitte geben Sie sinnvolle Stärken und Schwächen ein (mindestens 10 Zeichen, keine zufälligen Zeichen).")
        st.stop()
    if not name or not company:
        st.error("Bitte geben Sie Ihren Namen und den Unternehmensnamen ein.")
        st.stop()

    cv_text = ""
    job_text = ""

    # Lebenslauf verarbeiten
    if cv_file:
        if cv_file.type == "application/pdf":
            cv_text = extract_text_from_pdf(cv_file)
        elif cv_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            cv_text = extract_text_from_docx(cv_file)
    else:
        st.error("Bitte laden Sie einen Lebenslauf hoch.")
        st.stop()

    # Stellenprofil verarbeiten
    if job_file:
        job_text = extract_text_from_pdf(job_file)
    else:
        st.error("Bitte laden Sie ein Stellenprofil hoch.")
        st.stop()

    # Schlüsselwörter anzeigen
    if cv_text and job_text and not ("Fehler" in cv_text or "Fehler" in job_text):
        st.write("**Schlüsselwörter (Lebenslauf):**", ", ".join(extract_keywords(cv_text)))
        st.write("**Schlüsselwörter (Stellenprofil):**", ", ".join(extract_keywords(job_text)))

    # Motivationsschreiben generieren
    if cv_text and job_text:
        if "Fehler" in cv_text or "Fehler" in job_text:
            st.error("Fehler beim Verarbeiten der hochgeladenen Dateien.")
        else:
            cover_letter = generate_cover_letter(cv_text, job_text, strengths, weaknesses, name, company)
            st.write("**Motivationsschreiben:**")
            st.write(cover_letter)

            # Als DOCX speichern und Download anbieten
            filename = save_cover_letter(cover_letter)
            with open(filename, "rb") as file:
                st.download_button(
                    label="Motivationsschreiben herunterladen",
                    data=file,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.error("Bitte füllen Sie alle Felder aus.")